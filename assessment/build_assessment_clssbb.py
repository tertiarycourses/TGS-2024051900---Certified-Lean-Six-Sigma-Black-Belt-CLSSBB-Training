#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Rebuild the CLSSBB assessment set in the WSQ house format.

REFORMAT ONLY — the assessed content is carried over VERBATIM from the v1 papers
already registered on Drive:
    WA (SAQ) : 2 short-answer knowledge questions (K1, K2), 1 hour
    CS       : 1 scenario + 5 tasks (LO1-LO5 / A1-A5), 90 minutes
The number of questions, the instrument type, the wording and the criterion codes
are NOT changed — only the document structure is corrected.

What was wrong with v1 and is fixed here:
  * Questions started on PAGE 2, immediately under the instructions.
    -> Page 2 is now Trainee Information + FULL Instructions + Grading, nothing
       else; an explicit page break puts the questions on PAGE 3.
  * The assessor sign-off sat in a "For Official Use Only" block at the BACK of
    the paper (and was duplicated in the CS).
    -> Removed. The Grading block lives on page 2, per the house standard.
  * No boxed answer space for the candidate.
    -> Every question/task now carries a bordered answer box.
  * The submission note said "Submit your answers on the document provided".
    -> Now instructs upload to the LMS with a real clickable hyperlink.

Answer keys keep their existing model answers verbatim and omit Trainee
Information / Instructions / Grading, which belong to the candidate paper only.

Usage:  python3 assessment/build_assessment_clssbb.py
"""
import os
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)

# prodoc.py (WSQ cover page + page numbers) ships with the tertiary-lesson-plan skill.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand)
        break
import prodoc  # noqa: E402

# ─── course metadata ────────────────────────────────────────────────────────
TITLE       = "Certified Lean Six Sigma Black Belt (CLSSBB) Training"
COURSE_CODE = "TGS-2024051900"
Q_VER = A_VER = "v2"

OUT = HERE
ORG_LOGO = os.path.join(REPO, "courseware", "assets", "tertiary-infotech-logo.png")
if not os.path.exists(ORG_LOGO):
    ORG_LOGO = None
COURSE_LOGO = None

BRAND = RGBColor(0x1F, 0x6F, 0xEB)
DARK  = RGBColor(0x11, 0x18, 0x27)
GREY  = RGBColor(0x55, 0x5B, 0x66)

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"
FILL_GAP = 6

# ─── CONTENT — carried over VERBATIM from the registered v1 papers ──────────
# (criterion, context paragraph, question, model-answer bullet lines)
WRITTEN = [
    ("K1",
     "During a training session for the Certified Lean Six Sigma Black Belt (CLSSBB) course, "
     "participants are discussing the importance of process control performance metrics. They are "
     "particularly focused on how to select and analyze SPC charts to monitor process capabilities "
     "effectively.",
     "What are some key aspects of process control performance metrics that can be utilized in the "
     "Control Phase of Six Sigma?",
     ["Overview of Control Phase is essential for understanding process control.",
      "SPC Chart Selection and Analysis helps in monitoring process performance.",
      "Process Capabilities provide insights into the effectiveness of the processes."]),
    ("K2",
     "During a team meeting, the project manager discusses the need to address a recurring issue in "
     "the production line that is causing delays. The team decides to utilize the Analyze Phase of "
     "their Lean Six Sigma training to identify the root causes of the problem and implement "
     "effective problem-solving techniques.",
     "How can the Analyze Phase of Lean Six Sigma training assist the team in solving the production "
     "line issue?",
     ["The Analyze Phase helps in identifying the root causes of problems.",
      "It provides structured problem-solving techniques to address issues.",
      "Data analysis in this phase supports evidence-based decision making.",
      "Findings from the Analyze Phase guide the improvement actions that follow."]),
]

SCENARIO = (
    "In a mid-sized manufacturing company, production delays have been increasingly impacting "
    "customer satisfaction and overall profitability. The management team has identified that the "
    "assembly line for a key product is experiencing significant inefficiencies, leading to a "
    "backlog of orders. To address this challenge, the company has initiated a Six Sigma project "
    "aimed at reducing cycle time and improving throughput. The project team, led by a newly "
    "certified Black Belt, is tasked with defining clear project objectives that align with the "
    "company's performance standards. They conduct a thorough analysis of the current process, "
    "identifying bottlenecks and areas for improvement, and establish a project scope that includes "
    "a detailed timeline and resource allocation based on organizational requirements.\n\n"
    "As the project progresses, the team employs structured problem-solving techniques to execute "
    "the plan effectively. They utilize tools such as DMAIC (Define, Measure, Analyze, Improve, "
    "Control) to systematically address the issues identified. After implementing several process "
    "improvements, including the introduction of a new scheduling system and enhanced training for "
    "assembly line workers, the team evaluates the outcomes against the initial objectives. The "
    "results show a significant reduction in cycle time and an increase in production capacity. "
    "Based on these findings, the Black Belt recommends follow-up actions, including ongoing "
    "monitoring of key performance metrics and regular training sessions to sustain improvements. "
    "This structured approach not only resolves the immediate production issues but also sets a "
    "foundation for continuous improvement within the organization."
)

# (criterion, task, model-answer bullet lines)
CASE_TASKS = [
    ("LO1) (A1",
     "In the context of a mid-sized manufacturing company facing production delays, how should the "
     "newly certified Black Belt define Six Sigma project objectives to align with the company's "
     "performance standards and effectively address the identified inefficiencies in the assembly "
     "line?",
     ["The key problem is production delays that negatively impact customer satisfaction and "
      "profitability.",
      "Define clear, measurable project objectives tied to the company's performance standards.",
      "Objectives should target cycle time reduction and throughput improvement on the assembly line.",
      "Align the objectives with the identified bottlenecks so the project attacks the real constraint.",
      "Taught in: Lab 1 (Select and Charter Your Capstone Project) and Lab 3 (DMAIC at Depth - Y = f(X) and baseline sigma)."]),
    ("LO2) (A2",
     "In the context of a mid-sized manufacturing company facing production delays, how should the "
     "project team, led by a newly certified Black Belt, establish the scope of their Six Sigma "
     "project to effectively address the inefficiencies in the assembly line and align with "
     "organizational requirements?",
     ["Establish a scope that names the process boundaries — what is in and what is out.",
      "Include a detailed timeline for the project phases.",
      "Allocate resources based on organizational requirements and availability.",
      "Ensure the scope is achievable within the constraints and agreed with the sponsor.",
      "Taught in: Lab 1 (charter scope in/out) and Lab 6 (SIPOC, Process Mapping and Scope Governance)."]),
    ("LO3) (A3",
     "In the context of a mid-sized manufacturing company facing production delays, how should the "
     "project team execute the Six Sigma project using structured problem-solving techniques to "
     "ensure alignment with the project plan?",
     ["Execute using the DMAIC framework — Define, Measure, Analyze, Improve, Control.",
      "Apply structured problem-solving techniques at each phase rather than ad-hoc fixes.",
      "Systematically address the issues identified during the analysis.",
      "Monitor progress against the project plan and correct deviations early.",
      "Taught in: Lab 15 (variation, run charts, Pareto), Lab 16 (fishbone and 5 Whys) and Lab 17 (hypothesis testing)."]),
    ("LO4) (A4",
     "In the context of a Six Sigma project aimed at reducing cycle time and improving throughput in "
     "a mid-sized manufacturing company, how should the project team evaluate the outcomes to ensure "
     "that the improvements align with the initial project objectives?",
     ["Evaluate the outcomes against the objectives defined at the start of the project.",
      "Measure the reduction in cycle time and the increase in production capacity.",
      "Compare post-improvement performance with the baseline established in Measure.",
      "Confirm the improvements are statistically and practically significant, not chance variation.",
      "Taught in: Lab 26 (FMEA, pilot design, success criteria) and Lab 30 (financial benefits validation)."]),
    ("LO5) (A5",
     "In the context of the Six Sigma project implemented in a mid-sized manufacturing company, what "
     "follow-up actions should the project team recommend based on the process control performance "
     "metrics observed after the implementation of improvements?",
     ["Recommend ongoing monitoring of key performance metrics to sustain the gain.",
      "Schedule regular training sessions to maintain the improved way of working.",
      "Establish process controls so performance does not drift back after handover.",
      "Set a foundation for continuous improvement within the organization.",
      "Taught in: Lab 27 (SPC and chart selection), Lab 29 (control plan governance) and Lab 30 (handover and closure)."]),
]

# ─── house helpers (mirrors ~/.claude/skills/wsq-assessment/build_assessment.py) ──
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Arial"
    n.font.size = Pt(11)
    return doc


def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    return p


def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def answer_box(doc, lines=None, height_pt=150):
    """1x1 bordered box: `lines` renders the model answer; otherwise blank writing space."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True
        run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None)
            b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln)
            rr.font.size = Pt(10.5)
    else:
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt * 20)))
        trh.set(qn('w:hRule'), 'atLeast')
        trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_hyperlink(p, url, text):
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    p._p.append(link)
    return link


def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0


BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]


def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)


def grading(doc, what):
    """The assessor sign-off. Lives on PAGE 2 — never at the back of the paper."""
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0


def finish(doc, path):
    prodoc.add_page_numbers(doc)
    prodoc.enable_update_fields(doc)
    doc.save(path)
    print("  saved:", os.path.basename(path))


def _title_block(doc, subtitle):
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, subtitle, size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=12)


# ─── builders ───────────────────────────────────────────────────────────────
def build_wa(answers):
    doc = base_doc()
    kind = ("Written Assessment (SAQ) — Answer Key" if answers
            else "Written Assessment (SAQ)")
    prodoc.add_cover_page(doc, kind, TITLE, (A_VER if answers else Q_VER).lstrip("v"),
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO,
                          course_code=COURSE_CODE)
    _title_block(doc, "Answers to Written Assessment (SAQ)" if answers
                 else "Written Assessment (SAQ)")
    if not answers:
        # PAGE 2 — Trainee Information + full Instructions + Grading. Nothing else.
        candidate_block(doc)
        instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the "
                     "underpinning knowledge required for the course learning outcomes.")
        page_break(doc)          # -> questions start on PAGE 3
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge "
              "covered in the course slides.", size=10.5, italic=True, color=GREY, after=8)
    per_page = 1
    for i, (crit, ctx, q, lines) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=lines if answers else None)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    name = (f"Answer to WA (SAQ) - {TITLE}.docx" if answers else f"WA (SAQ) - {TITLE}.docx")
    finish(doc, os.path.join(OUT, name))


def build_cs(answers):
    doc = base_doc()
    kind = "Case Study (CS) — Answer Key" if answers else "Case Study (CS)"
    prodoc.add_cover_page(doc, kind, TITLE, (A_VER if answers else Q_VER).lstrip("v"),
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO,
                          course_code=COURSE_CODE)
    _title_block(doc, "Answers to Case Study Assessment" if answers
                 else "Case Study Assessment")
    if not answers:
        # PAGE 2 — Trainee Information + full Instructions + Grading. Nothing else.
        candidate_block(doc)
        instructions(doc, "90 minutes")
        grading(doc, "Candidate has completed all case study tasks and can justify the approach "
                     "taken against the course learning outcomes.")
        page_break(doc)          # -> scenario and tasks start on PAGE 3
    para(doc, "Case Study", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    for chunk in SCENARIO.split("\n\n"):
        para(doc, chunk, size=11, after=8)
    para(doc, "Answer all tasks with reference to the scenario above.",
         size=10.5, italic=True, color=GREY, after=8)
    # NO page break here — the house rule is that the scenario AND the tasks begin
    # on page 3. Breaking after the scenario pushes Task 1 to page 4 and fails the
    # three-page rule (caught by verify_assessment.py).
    for i, (crit, task, lines) in enumerate(CASE_TASKS, 1):
        para(doc, f"Task {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, f"{task}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=lines if answers else None)
        if i < len(CASE_TASKS):
            page_break(doc)
    name = (f"Answer to CS - {TITLE}.docx" if answers else f"CS - {TITLE}.docx")
    finish(doc, os.path.join(OUT, name))


if __name__ == "__main__":
    print("Building CLSSBB assessment set (house format):")
    build_wa(False)
    build_wa(True)
    build_cs(False)
    build_cs(True)
    print("Done. 4 documents in", OUT)
