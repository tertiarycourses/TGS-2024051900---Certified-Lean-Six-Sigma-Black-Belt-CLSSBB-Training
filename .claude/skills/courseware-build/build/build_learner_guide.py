#!/usr/bin/env python3
"""Generate the CLSSBB Learner Guide as BOTH a Markdown mirror (LG-*.md at repo
root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per DMAIC phase, one sub-section per lab (Objective · Goal ·
What you'll build · Step-by-step · Check your work), plus quick-reference
formulas, assessment preparation and a glossary. All content is driven by
course_data + the domain data files, keeping the LG 100% aligned with the slide
deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
from data_domain6 import DOMAIN6
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5+DOMAIN6
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")
# Lab visuals are generated into the COURSE repo (courseware/assets/lg-visuals),
# which is a different place from the skill-local ASSETS dir used for the logos.
REPO_ASSETS=os.path.join(REPO,"courseware","assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def img(path,caption): B.append(("img",path,caption))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It follows the DMAIC roadmap end to end — Define, Measure, Analyze, Improve, Control — and provides "
  "step-by-step instructions for every hands-on lab. Every lab is a capstone building block - its "
  "output goes directly into the project you present and defend on Day 5.")
p("The course content is grounded in the body of knowledge published by The Council for Six Sigma "
  "Certification (CSSC) in 'Six Sigma: A Complete Step-by-Step Guide', so what you learn here matches "
  "the recognised Black Belt standard.")
p("This course is the deliberate step up from the Green Belt (CLSSGB). DMAIC is recapped "
  "IN DEPTH — because a Black Belt must be able to MENTOR Green Belts through these tools, and you cannot "
  "coach what you only half-recall — and then extended with the advanced material the Green Belt course "
  "does not teach: attribute MSA and Kappa analysis, nested and destructive gauge studies, multiple "
  "regression with model diagnostics, ANOVA and interaction effects, non-parametric tests, multi-vari "
  "studies, full and fractional factorial Design of Experiments, response surface methodology, Taguchi "
  "robust design, CUSUM, EWMA and multivariate SPC, change leadership, and financial benefits validation.")
p("YOUR CAPSTONE PROJECT RUNS FROM DAY 1. In Lab 1 you select a real process from your own workplace and "
  "charter it. Every lab thereafter adds one artifact to that same project — the measurement study, the "
  "regression model, the experiment design, the control plan. On Day 5 you consolidate those artifacts "
  "into an A3 storyboard and present and defend it to a steering committee. The capstone is an "
  "ADDITIONAL deliverable on top of the WSQ assessment (Written Assessment SAQ + Case Study), which "
  "you also sit on Day 5. A worked reference scenario (Meridian Medical Devices — infusion pump assemblies across "
  "three plants) is provided throughout for learners who cannot use workplace data.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework Alignment")
p(f"This course is aligned to the WSQ Technical Skills and Competencies (TSC) {C.TSC_TITLE} ({C.TSC_CODE}).")
h3("Abilities")
bullets(C.TSC_ABILITIES)
h3("Knowledge")
bullets(C.TSC_KNOWLEDGE)

h1("Before You Start")
h3("What you need")
bullets([
 "A laptop with a spreadsheet application (Microsoft Excel, Google Sheets or LibreOffice Calc).",
 "A browser, for the interactive problem-solving tools used in the Analyze labs.",
 "The course slides and this Learner Guide, downloaded from https://lms-tms.tertiaryinfotech.com.",
 "A REAL process from your own workplace, with accessible data — this becomes your capstone project.",
])
h3("The interactive problem-solving toolkit")
p("Five browser-based tools are used during the labs. No installation or licence is required.")
bullets([
 "5 Whys — build and share a 5 Whys chain: https://alfredang.github.io/5whys/",
 "Fishbone Diagram — build an Ishikawa cause-and-effect diagram: https://alfredang.github.io/fishbone/",
 "Pareto Chart (collaborative) — your team brainstorms and votes in one live session and the Pareto chart builds itself: https://alfredang.github.io/paretochart/",
 "NovaSPC — run charts, SPC charts and process capability from your own CSV: https://alfredang.github.io/novaspc/",
 "SIPOC — build a SIPOC map (Suppliers, Inputs, Process, Outputs, Customers) and agree the process boundaries: https://alfredang.github.io/sipoc/",
])
h3("How the labs work")
bullets([
 "Every lab is a CAPSTONE BUILDING BLOCK — its output goes straight into your capstone project pack.",
 "Labs are completed in order because later labs consume the charter, baseline, measurement study and analysis built earlier.",
 "Apply every lab to your OWN workplace process; use the Meridian Medical Devices scenario only as a fallback.",
])
h3("Conventions used in every lab")
bullets([
 "Each lab states its objective, the deliverable you produce, the steps, and a check to confirm you are done.",
 "Tables shown in the steps can be built in a spreadsheet or on the worksheet provided.",
 "Where a lab uses an online tool, the tool URL is shown with the step.",
 "Keep every lab output — they combine into the capstone project you present and defend on Day 5.",
])

# ---------------- per-topic, per-lab ----------------
for t in C.TOPICS:
    label = t["phase"].title() if t["num"] else "Foundations"
    h1(f"{t['phase']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets([f"{name} — {desc}" for name, desc in t["concepts"]])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        title = a["title"].replace("Elective — ","")
        h2(f"Lab {a['num']} — {title}")
        p(f"Objective: {a['objective']}")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools and techniques: {a['services']}.)")
        _vis=os.path.join(REPO_ASSETS,"lg-visuals",f"lab-{a['num']:02d}-visual.png")
        if os.path.exists(_vis):
            img(_vis,f"Lab {a['num']} at a glance — the deliverable, the tools and the steps.")
        h3("Step-by-step")
        steps([(instr,cmd) for instr,cmd in a["steps"]])
        h3("Check your work")
        p(a["test"])
        note(f"The full worksheet for this lab is in labs/lab-{a['num']:02d}-*.md.")
        rule()

h1("Quick Reference — Formulas You Should Know")
h3("Process performance metrics")
bullets([
 "Yield = (Good units / Total units) x 100",
 "DPU (Defects Per Unit) = Defects / Units",
 "DPO (Defects Per Opportunity) = Defects / (Units x Opportunities per unit)",
 "DPMO (Defects Per Million Opportunities) = DPO x 1,000,000",
 "First Pass Yield (FPY) = units passing with no rework / total units started",
 "Rolled Throughput Yield (RTY) = FPY(step 1) x FPY(step 2) x ... x FPY(step n)",
 "Process Cycle Efficiency = Value-added time / Total lead time",
 "Takt time = Available working time / Customer demand",
])
h3("Sigma level reference")
bullets([
 "1 sigma — 690,000 DPMO (about 31% yield)",
 "2 sigma — 308,000 DPMO (about 69% yield)",
 "3 sigma — 66,800 DPMO (about 93.3% yield)",
 "4 sigma — 6,210 DPMO (about 99.38% yield)",
 "5 sigma — 233 DPMO (about 99.977% yield)",
 "6 sigma — 3.4 DPMO (about 99.99966% yield)",
])

h1("Quick Reference — The Eight Wastes (DOWNTIME)")
bullets([
 "D — Defects: output that fails the requirement and must be corrected or redone.",
 "O — Overproduction: producing more, or earlier, than the customer needs.",
 "W — Waiting: work or people idle, waiting for the next step, an approval or information.",
 "N — Non-utilised talent: skills and ideas of people not being used.",
 "T — Transport: unnecessary movement of materials, work items or information between places.",
 "I — Inventory: work in progress, backlogs and queues sitting between steps.",
 "M — Motion: unnecessary movement of people, or switching between systems and screens.",
 "E — Extra-processing: doing more work to the output than the customer requires or values.",
])

h1("Preparing for the Assessment")
bullets([
 C.ASSESSMENT["written"],
 C.ASSESSMENT["practical"],
 C.ASSESSMENT["capstone"],
 "The WSQ assessment is the Written Assessment (SAQ) plus the Case Study (CS) — both open book.",
 "Your capstone is ADDITIONAL to the WSQ papers; it is built from the artifacts of Labs 1 to 30.",
 "Revise by re-working your own lab outputs — they cover every tool the papers can ask about.",
 "Open book: you may use these slides, this Learner Guide and your own project data throughout.",
 "Consolidate your artifacts into a ONE-PAGE A3 storyboard: Background, Current State, Goal, Root Cause, Countermeasures, Results, Follow-up.",
 "Lead your presentation with the decision, the benefit and the ask — the sponsor needs those in the first two minutes.",
 "Prepare an APPENDIX with your full analysis; this is where challenge questions get answered with evidence.",
 "Be ready to defend your sample size and the statistical power of your tests.",
 "Be ready to prove your measurement system is valid — Gage R&R for continuous data, Kappa for attribute data.",
 "Be ready to state which assumptions you tested (normality, equal variance, independence) and what you did when they failed.",
 "Be ready to answer a confounding challenge: what else changed, and why is the gain attributable to your intervention?",
 "Be ready to state your DOE resolution and exactly which effects were aliased.",
 "Be ready to separate hard, soft and cost-avoidance benefits, and name who in Finance validated them.",
 "Be ready to explain what stops the improvement decaying once attention moves elsewhere.",
 "Where a challenge exposes a genuine weakness, acknowledge it and state the mitigation — that is stronger than defending it.",
 C.ASSESSMENT["note"],
])

h1("Glossary")
gl=[
 ("Lean","A method to maximise customer value by systematically identifying and removing waste."),
 ("Six Sigma","A data-driven method to reduce variation and defects; the target is 3.4 defects per million opportunities."),
 ("Lean Six Sigma","The combined method — Lean improves speed and flow, Six Sigma improves consistency and accuracy."),
 ("DMAIC","Define, Measure, Analyze, Improve, Control — the Six Sigma improvement roadmap."),
 ("PDCA","Plan, Do, Check, Act — a lighter improvement cycle used for small, fast improvements."),
 ("VOC","Voice of the Customer — customer needs and expectations expressed in the customer's own words."),
 ("CTQ","Critical to Quality — a specific, measurable requirement translated from a VOC statement."),
 ("SIPOC","Suppliers, Inputs, Process, Outputs, Customers — the macro 'as-is' process map."),
 ("Muda","The Japanese term for waste — any activity that consumes resource but creates no customer value."),
 ("DOWNTIME","Mnemonic for the eight wastes: Defects, Overproduction, Waiting, Non-utilised talent, Transport, Inventory, Motion, Extra-processing."),
 ("Value-added (VA)","An activity the customer would be willing to pay for because it changes the product or service."),
 ("Non-value-added (NVA)","An activity that consumes resource but adds nothing the customer values — pure waste."),
 ("Defect","An output that fails to meet the CTQ requirement."),
 ("Common cause variation","Natural, always-present variation built into the process; addressed by changing the process."),
 ("Special cause variation","Unusual variation traceable to a specific assignable event; addressed by investigating that event."),
 ("Pareto principle","The 80/20 rule — roughly 80% of effects come from 20% of causes."),
 ("Run chart","A plot of a metric over time, used to spot trend, shift, cluster and oscillation patterns."),
 ("5 Whys","Repeatedly asking 'why' to drill from a symptom down to an actionable root cause."),
 ("Fishbone (Ishikawa) diagram","A cause-and-effect diagram organising candidate causes into categories such as the 5Ms."),
 ("5M categories","Manpower, Method, Machine, Material, Measurement — standard Fishbone categories."),
 ("Multi-voting","A team technique to narrow a long list of candidate causes to an agreed shortlist."),
 ("DPU / DPO / DPMO","Defects per unit / per opportunity / per million opportunities — standard defect-rate metrics."),
 ("Yield","The percentage of output produced without defects."),
 ("RTY","Rolled Throughput Yield — the probability a unit passes through every process step with no rework."),
 ("Sigma level","A common yardstick for process performance, derived from DPMO."),
 ("COPQ","Cost of Poor Quality — the internal failure, external failure, appraisal and prevention costs of defects."),
 ("5S","Sort, Set in order, Shine, Standardise, Sustain — a Lean method for organising the workplace."),
 ("Poka-Yoke","Mistake proofing — designing the process so an error is difficult or impossible to make."),
 ("Standard work","The documented current best-known method for performing a task."),
 ("Kaizen","Continuous improvement through many small changes, made by everyone."),
 ("Takt time","Available working time divided by customer demand — the pace the process must run to meet demand."),
 ("Control plan","The document naming the metric, target, monitoring method, frequency, owner and reaction plan."),
 ("Specification limits (LSL/USL)","The acceptable limits set by the customer; outside them is a defect."),
 ("Control limits (LCL/UCL)","Limits calculated from the process itself, typically the mean plus/minus three standard deviations."),
 ("SPC","Statistical Process Control — monitoring a process over time to predict and prevent problems."),
 ("A3","A one-page report telling the whole improvement story from background through to follow-up."),
 ("Gemba","'The real place' — going to where the work actually happens to observe the process directly."),
]
B.append(("dl",gl))


# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="img":
            _p,_cap=rest[0],rest[1]
            _rel=os.path.relpath(_p,REPO)
            out+=[f"![{_cap}]({_rel})","",f"*{_cap}*",""]
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("1","1 June 2026","Initial release - CLSSBB Learner Guide.",C.TRAINER),
 ("2",C.VERSION_DATE,"Major revision: rebuilt as a 5-day Black Belt progression from the Green Belt "
  "course (CLSSGB). DMAIC is recapped IN DEPTH and then extended with Black-Belt-only content: "
  "attribute MSA and Kappa analysis, nested and destructive gauge studies, non-normal capability and "
  "Cpk versus Ppk, multiple regression with model diagnostics and VIF, ANOVA with interaction effects, "
  "non-parametric tests, multi-vari studies, full and fractional factorial Design of Experiments with "
  "confounding and resolution, response surface methodology, Taguchi robust design, CUSUM, EWMA, "
  "short-run and multivariate SPC, change leadership and resistance management, financial benefits "
  "validation, and mentoring Green Belts. Expanded to 34 hands-on labs across five days. THE CAPSTONE "
  "PROJECT RUNS FROM DAY 1 - each learner charters a real workplace process in Lab 1 and every "
  "subsequent lab adds one artifact to it; Day 5 consolidates them into an A3 storyboard and a steering "
  "committee presentation.",C.TRAINER),
 ("3",C.VERSION_DATE,"Corrected the assessment model to match the registered WSQ instrument: the assessment is the Written Assessment (WA) Short-Answer Questions plus the Case Study (CS), both open book. The capstone project and steering committee presentation are retained as an ADDITIONAL course deliverable on top of the WSQ assessment, not as a replacement for it. Added the online SIPOC tool (alfredang.github.io/sipoc) to the toolkit, Lab 6 and the Define slides.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        # Write the step number as literal text rather than using Word's
        # "List Number" style: that style shares ONE counter for the whole
        # document, so Lab 9 would continue from Lab 8 (…62, 63, 64) instead of
        # restarting at 1. An explicit number restarts correctly for every lab.
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph()
            para.paragraph_format.left_indent=Pt(18)
            para.paragraph_format.first_line_indent=Pt(-18)
            r=para.add_run(f"{i}.  "); r.bold=True; r.font.color.rgb=BRAND
            para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="img":
        _p,_cap=rest[0],rest[1]
        try:
            doc.add_picture(_p,width=Inches(6.2))
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=cp.add_run(_cap); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GREY
        except Exception:
            pass
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
